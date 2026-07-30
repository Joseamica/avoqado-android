#!/usr/bin/env python3
"""Actualiza la guía Word conservando su diseño y listas OOXML."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def paragraph_with_text(document: Document, text: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise ValueError(f"Se esperaba un párrafo para {text!r}; encontrados: {len(matches)}")
    return matches[0]


def replace_text(paragraph: Paragraph, text: str) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.add_run(text)


def insert_clone_after(paragraph: Paragraph, template: Paragraph, text: str) -> Paragraph:
    element = copy.deepcopy(template._p)
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    replace_text(inserted, text)
    return inserted


def create_numbering_instance(document: Document, template: Paragraph) -> int:
    template_num = template._p.pPr.numPr
    template_num_id = int(template_num.numId.val)
    numbering = document.part.numbering_part.element
    source_num = numbering.xpath(f"./w:num[@w:numId='{template_num_id}']")[0]
    abstract_num_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    number_properties = properties.get_or_add_numPr()
    number_properties.get_or_add_ilvl().val = 0
    number_properties.get_or_add_numId().val = num_id


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    document = Document(args.input)

    version_cell = document.tables[0].cell(0, 1)
    version_cell.paragraphs[1].text = "1.1 · 30 de julio de 2026"

    replace_text(
        paragraph_with_text(document, "captura el peso en Peso (kg), por ejemplo 0.435;"),
        "si aparece Báscula, espera el estado estable; Avoqado llenará el peso;",
    )
    replace_text(
        paragraph_with_text(document, "revisa el precio por kilogramo y el total mostrado;"),
        "si no está conectada, pulsa Capturar manualmente y escribe el peso visible en Peso (kg);",
    )
    replace_text(
        paragraph_with_text(document, "pulsa Agregar • $…; el botón muestra el total calculado."),
        "revisa el precio por kilogramo y el total, y pulsa Agregar • $…",
    )

    heading = paragraph_with_text(document, "Operación disponible desde el primer día")
    old_intro = paragraph_with_text(
        document,
        "El flujo de vales funciona aunque una báscula no esté conectada a Avoqado. Mientras se certifica la lectura automática de cada modelo, el operador captura en la app el peso que muestra la báscula.",
    )
    old_rhino = paragraph_with_text(
        document,
        "Cremería — Rhino: pesar el producto, esperar una lectura estable y capturar los kilogramos en Avoqado.",
    )
    old_justa = paragraph_with_text(
        document,
        "CEDIS — Justa LP7516: usar para recepción, salida o conteo de inventario según el proceso configurado. No genera por sí sola un vale de venta de cremería.",
    )
    old_fallback = paragraph_with_text(
        document,
        "Cuando la lectura automática sea habilitada para una báscula, el operador seguirá revisando tres cosas antes de aceptar el peso: cero o tara correcta, lectura estable y unidad en kilogramos. Si la lectura automática falla, se continúa con captura manual; caja y los vales no se detienen.",
    )
    numbered_template = paragraph_with_text(document, "Abre Menú → Productos.")
    plain_template = paragraph_with_text(
        document,
        "Después de configurarlo, tocar ese producto desde Cobrar abrirá directamente el panel Peso. El teclado grande que aparece en la pestaña Teclado sirve para importes libres; no debe usarse para pesar.",
    )

    replace_text(heading, "Cremería — Rhino BAR-8RS")
    rhino_num_id = create_numbering_instance(document, numbered_template)
    rhino_steps = [
        (old_intro, "Entra a Cobrar y toca un producto configurado como Se vende por peso."),
        (old_rhino, "Coloca el producto y espera a que la tarjeta de báscula muestre una lectura estable."),
        (old_justa, "Revisa cero o tara, kilogramos y total."),
        (old_fallback, "Pulsa Agregar. Si el cable o la lectura fallan, usa Capturar manualmente."),
    ]
    for paragraph, text in rhino_steps:
        replace_text(paragraph, text)
        apply_numbering(paragraph, rhino_num_id)

    cursor = old_fallback
    cedis_heading = insert_clone_after(cursor, heading, "CEDIS — Justa LP7516")
    cursor = cedis_heading
    cedis_num_id = create_numbering_instance(document, numbered_template)
    for text in [
        "Entra a Inventario y abre Conteos.",
        "Inicia un conteo completo o cíclico.",
        "Selecciona un insumo cuya unidad sea kilogramo o gramo.",
        "Coloca el producto en la Justa y espera una lectura estable.",
        "Revisa el valor y pulsa Usar este peso. Avoqado llena el campo, pero no guarda sin tu confirmación.",
        "Continúa con Siguiente artículo y termina con Revisar conteo.",
    ]:
        cursor = insert_clone_after(cursor, numbered_template, text)
        apply_numbering(cursor, cedis_num_id)

    cursor = insert_clone_after(
        cursor,
        plain_template,
        "Para artículos en piezas, litros u otras unidades se usa el teclado normal. La recepción de órdenes de compra permanece manual en esta etapa.",
    )
    insert_clone_after(
        cursor,
        plain_template,
        "En ambos lugares revisa cero o tara, lectura estable y unidad. Si la lectura automática falla, captura manualmente el peso visible; caja, vales e inventario no se detienen.",
    )

    replace_text(
        paragraph_with_text(
            document,
            "Cremería pesa el jamón, captura 0.250 kg, emite su vale y conserva el jamón.",
        ),
        "Cremería pesa el jamón, usa la lectura estable o captura 0.250 kg, emite su vale y conserva el jamón.",
    )

    args.out.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.out)


if __name__ == "__main__":
    main()
